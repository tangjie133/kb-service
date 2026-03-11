#!/usr/bin/env python3
"""Document processing and chunking for multiple formats"""

import re
from typing import List, Dict, Optional
from dataclasses import dataclass
from pathlib import Path
import yaml
import logging

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Document chunk with metadata"""
    id: str
    content: str
    metadata: Dict
    source_file: str
    start_line: int
    end_line: int


class DocumentProcessor:
    """Process documents into chunks (supports .md, .txt, .pdf, .docx)"""
    
    SUPPORTED_EXTENSIONS = {'.md', '.txt', '.pdf', '.docx'}
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_file(self, file_path: Path) -> List[DocumentChunk]:
        """
        Process a single file based on its extension
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.warning(f"File not found: {file_path}")
            return []
        
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"Unsupported file type: {extension}")
            return []
        
        try:
            if extension == '.md':
                return self._process_markdown(file_path)
            elif extension == '.txt':
                return self._process_text(file_path)
            elif extension == '.pdf':
                return self._process_pdf(file_path)
            elif extension == '.docx':
                return self._process_docx(file_path)
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return []
        
        return []
    
    def _process_markdown(self, file_path: Path) -> List[DocumentChunk]:
        """Process Markdown files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        metadata = self._extract_metadata(content)
        metadata['filename'] = file_path.name
        metadata['file_type'] = 'markdown'
        
        body = self._remove_metadata(content)
        chunks = self._split_text(body)
        
        return self._create_chunks(file_path, chunks, metadata, body)
    
    def _process_text(self, file_path: Path) -> List[DocumentChunk]:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        metadata = {
            'filename': file_path.name,
            'file_type': 'text',
            'title': file_path.stem
        }
        
        chunks = self._split_text(content)
        return self._create_chunks(file_path, chunks, metadata, content)
    
    def _process_pdf(self, file_path: Path) -> List[DocumentChunk]:
        """Process PDF files"""
        try:
            import PyPDF2
        except ImportError:
            logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
            return []
        
        content = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    content += f"\n--- Page {page_num + 1} ---\n"
                    content += page.extract_text() or ""
        
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return []
        
        metadata = {
            'filename': file_path.name,
            'file_type': 'pdf',
            'title': file_path.stem,
            'pages': num_pages
        }
        
        chunks = self._split_text(content)
        return self._create_chunks(file_path, chunks, metadata, content)
    
    def _process_docx(self, file_path: Path) -> List[DocumentChunk]:
        """Process Word documents"""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return []
        
        try:
            doc = Document(file_path)
            
            # Extract text
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                content_parts.append("\n[Table]")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content_parts.append(row_text)
            
            content = "\n\n".join(content_parts)
            
            # Try to get title from first paragraph
            title = file_path.stem
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                title = doc.paragraphs[0].text.strip()[:100]
        
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return []
        
        metadata = {
            'filename': file_path.name,
            'file_type': 'docx',
            'title': title
        }
        
        chunks = self._split_text(content)
        return self._create_chunks(file_path, chunks, metadata, content)
    
    def _create_chunks(
        self, 
        file_path: Path, 
        chunks: List[str], 
        metadata: Dict,
        full_content: str
    ) -> List[DocumentChunk]:
        """Create DocumentChunk objects from text chunks"""
        doc_chunks = []
        
        for i, chunk_content in enumerate(chunks):
            chunk_id = f"{file_path.stem}_{i}"
            
            # Calculate line numbers
            start_pos = full_content.find(chunk_content)
            if start_pos >= 0:
                start_line = full_content[:start_pos].count('\n') + 1
                end_line = start_line + chunk_content.count('\n')
            else:
                start_line = 0
                end_line = 0
            
            doc_chunks.append(DocumentChunk(
                id=chunk_id,
                content=chunk_content.strip(),
                metadata=metadata.copy(),
                source_file=str(file_path),
                start_line=start_line,
                end_line=end_line
            ))
        
        return doc_chunks
    
    def _extract_metadata(self, content: str) -> Dict:
        """Extract YAML frontmatter and other metadata"""
        metadata = {}
        
        # YAML frontmatter
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                try:
                    metadata = yaml.safe_load(parts[1]) or {}
                except yaml.YAMLError as e:
                    logger.warning(f"YAML parsing error: {e}")
        
        # Extract title from first H1
        if 'title' not in metadata:
            title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
            if title_match:
                metadata['title'] = title_match.group(1).strip()
        
        # Extract tags
        if 'tags' not in metadata:
            tags_match = re.search(r'[Tt]ags?:\s*\[?([^\]]+)\]?', content)
            if tags_match:
                tags_str = tags_match.group(1)
                metadata['tags'] = [t.strip() for t in tags_str.split(',')]
        
        return metadata
    
    def _remove_metadata(self, content: str) -> str:
        """Remove YAML frontmatter from content"""
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                return parts[2].strip()
        return content
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks
        Strategy: Split by paragraphs, then combine
        """
        if not text.strip():
            return []
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Split by double newlines (paragraphs)
        paragraphs = re.split(r'\n\n+', text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        if not paragraphs:
            return []
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for para in paragraphs:
            para_length = len(para)
            
            # If adding this paragraph exceeds chunk size
            if current_length + para_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append('\n\n'.join(current_chunk))
                
                # Keep overlap
                overlap_chunks = []
                overlap_length = 0
                for p in reversed(current_chunk):
                    if overlap_length + len(p) > self.chunk_overlap:
                        break
                    overlap_chunks.insert(0, p)
                    overlap_length += len(p)
                
                current_chunk = overlap_chunks + [para]
                current_length = overlap_length + para_length
            else:
                current_chunk.append(para)
                current_length += para_length + 2  # +2 for \n\n
        
        # Don't forget the last chunk
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks
    
    def get_supported_files(self, directory: Path) -> List[Path]:
        """Get all supported files in a directory"""
        files = []
        for ext in self.SUPPORTED_EXTENSIONS:
            files.extend(directory.rglob(f"*{ext}"))
        return files
